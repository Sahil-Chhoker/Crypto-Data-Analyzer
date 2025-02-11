import requests
import pandas as pd
import time
from datetime import datetime
from openpyxl.styles import PatternFill, Font
import threading
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns


# Main tracker class
class CryptoTracker:
    def __init__(self, update_interval=300):  # 5 min default interval
        # API endpoints - using multiple for redundancy (source https://developers.binance.com/docs/binance-spot-api-docs/rest-api/general-api-information)
        self.endpoints = [
            "https://api.binance.com",
            "https://api-gcp.binance.com",
            "https://api1.binance.com",
        ]
        
        # For fetching market data
        self.data_api = "https://data-api.binance.vision"
        self.coingecko = "https://api.coingecko.com/api/v3"
        
        self.update_interval = update_interval
        self.last_report_hour = None
        
        # Output files
        self.excel_file = "crypto_data.xlsx"
        self.report_file = "market_report.docx"
        
        self.running = True

    def get_coin_names(self):
        try:
            r = requests.get(f"{self.coingecko}/coins/list")
            
            if r.status_code == 200:
                coins = r.json()
                return {c['symbol'].upper(): c['name'] for c in coins}
            
            print("Couldn't fetch coin names")
            return {}
            
        except Exception as e:
            print(f"Error getting coin names: {e}")
            return {}

    def get_market_data(self):
        try:
            r = requests.get(f"{self.data_api}/api/v3/ticker/24hr")
            
            if r.status_code != 200:
                print("API call failed")
                return None
                
            data = r.json()
            
            usdt_pairs = [x for x in data if x['symbol'].endswith('USDT')]
            
            names = self.get_coin_names()
            for pair in usdt_pairs:
                sym = pair['symbol'].replace('USDT', '').upper()
                pair['name'] = names.get(sym, sym)  # Use symbol if name not found
            
            return usdt_pairs
            
        except Exception as e:
            print(f"Failed to get market data: {e}")
            return None

    def process_data(self, raw_data):
        df = pd.DataFrame(raw_data)
        
        # Convert strings to numbers
        number_cols = [
            'lastPrice', 'volume', 'quoteVolume', 
            'priceChangePercent', 'weightedAvgPrice'
        ]
        
        for col in number_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # Sort by volume and get top 50
        df = df.sort_values('quoteVolume', ascending=False).head(50)
        
        # Calculate market cap (price * volume)
        df['market_cap'] = df['lastPrice'] * df['volume']
        
        final = pd.DataFrame({
            'Name': df['name'],
            'Symbol': df['symbol'].str.replace('USDT', ''),
            'Price': df['lastPrice'],
            'Market Cap': df['market_cap'],
            'Volume (24h)': df['volume'],
            'Change (24h)': df['priceChangePercent'],
            'Avg Price': df['weightedAvgPrice'],
        })
        
        return final

    def make_chart(self, df):
        plt.figure(figsize=(12, 6))
        
        # Distribution of 24h price changes
        sns.histplot(data=df, x='Change (24h)', bins=20)
        plt.title('24h Price Changes')
        plt.xlabel('Change (%)')
        plt.ylabel('Number of Coins')
        
        plt.savefig('price_changes.png')
        plt.close()

    def analyze_market(self, df):
        # Calculate market overview statistics
        stats = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'market_overview': {
                'total_market_cap': df['Market Cap'].sum(),
                'total_volume_24h': df['Volume (24h)'].sum(),
                'avg_price': df['Price'].mean(),
                'median_price': df['Price'].median(),
            },
            'market_health': {
                'positive_performers': len(df[df['Change (24h)'] > 0]),
                'negative_performers': len(df[df['Change (24h)'] < 0]),
            },
            'detailed_analysis': {
                'top_5_by_market_cap': df.nlargest(5, 'Market Cap')[['Name', 'Symbol', 'Market Cap', 'Price']].to_dict('records'),
                'price_statistics': {
                    'average_price': df['Price'].mean(),
                    'median_price': df['Price'].median(),
                    'highest_price': df['Price'].max(),
                    'lowest_price': df['Price'].min(),
                },
                'price_changes': {
                    'highest_24h_change': df.nlargest(1, 'Change (24h)')[['Name', 'Symbol', 'Change (24h)']].to_dict('records')[0],
                    'lowest_24h_change': df.nsmallest(1, 'Change (24h)')[['Name', 'Symbol', 'Change (24h)']].to_dict('records')[0],
                    'average_24h_change': df['Change (24h)'].mean(),
                }
            }
        }
        return stats

    def make_report(self, df, stats):
        doc = Document()
        
        # Title
        doc.add_heading('Crypto Market Report', 0)
        doc.add_paragraph(f"Generated: {stats['timestamp']}")
        
        # Market overview section
        doc.add_heading('Market Overview', 1)
        overview = stats['market_overview']
        doc.add_paragraph(f"Total Market Cap: ${overview['total_market_cap']:,.2f}")
        doc.add_paragraph(f"24h Volume: ${overview['total_volume_24h']:,.2f}")
        
        # Top coins section with more details
        doc.add_heading('Top 5 Cryptocurrencies by Market Cap', 1)
        for coin in stats['detailed_analysis']['top_5_by_market_cap']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{coin['Name']} ({coin['Symbol']})\n").bold = True
            p.add_run(f"Market Cap: ${coin['Market Cap']:,.2f}\n")
            p.add_run(f"Price: ${coin['Price']:,.2f}")
        
        # Price Statistics
        doc.add_heading('Price Statistics', 1)
        price_stats = stats['detailed_analysis']['price_statistics']
        doc.add_paragraph(f"Average Price: ${price_stats['average_price']:,.2f}")
        doc.add_paragraph(f"Median Price: ${price_stats['median_price']:,.2f}")
        doc.add_paragraph(f"Highest Price: ${price_stats['highest_price']:,.2f}")
        doc.add_paragraph(f"Lowest Price: ${price_stats['lowest_price']:,.2f}")
        
        # 24h Price Changes
        doc.add_heading('24-Hour Price Changes', 1)
        changes = stats['detailed_analysis']['price_changes']
        best = changes['highest_24h_change']
        worst = changes['lowest_24h_change']
        
        doc.add_paragraph(f"Highest Gainer: {best['Name']} ({best['Symbol']}) with {best['Change (24h)']:+.2f}%")
        doc.add_paragraph(f"Biggest Decliner: {worst['Name']} ({worst['Symbol']}) with {worst['Change (24h)']:+.2f}%")
        doc.add_paragraph(f"Average 24h Change: {changes['average_24h_change']:+.2f}%")
        
        # Market health section
        doc.add_heading('Market Health', 1)
        health = stats['market_health']
        doc.add_paragraph(f"Coins Up: {health['positive_performers']}")
        doc.add_paragraph(f"Coins Down: {health['negative_performers']}")
        
        # Add a price change chart
        self.make_chart(df)
        doc.add_picture('price_changes.png', width=Inches(6))
        
        doc.save(self.report_file)

    def update_excel(self, df, stats):
        with pd.ExcelWriter(self.excel_file, engine='openpyxl') as writer:
            # Main data sheet
            df.to_excel(writer, sheet_name='Market Data', index=False)
            
            # Analysis sheet
            analysis = [
                ['Last Updated', stats['timestamp']],
                ['Market Overview', ''],
                ['Total Market Cap', f"${stats['market_overview']['total_market_cap']:,.2f}"],
                ['24h Volume', f"${stats['market_overview']['total_volume_24h']:,.2f}"],
                ['Average Price', f"${stats['market_overview']['avg_price']:,.2f}"],
                ['', ''],
                ['Market Health', ''],
                ['Coins Up', stats['market_health']['positive_performers']],
                ['Coins Down', stats['market_health']['negative_performers']],
            ]
            
            pd.DataFrame(analysis).to_excel(
                writer, sheet_name='Analysis',
                index=False, header=False
            )
            
            # Make it look nice
            self._format_excel(writer)

    # Add some color to the Excel sheets
    def _format_excel(self, writer):
        for sheet in ['Market Data', 'Analysis']:
            ws = writer.sheets[sheet]
            # Color the headers
            for cell in ws[1]:
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)

    # Main loop
    def run(self):
        while self.running:
            print(f"\nUpdating data... {datetime.now().strftime('%H:%M:%S')}")
            
            try:
                # Get fresh data
                data = self.get_market_data()
                if not data:
                    continue
                    
                # Process it
                df = self.process_data(data)
                stats = self.analyze_market(df)
                
                # Update files
                self.update_excel(df, stats)
                
                # Make report on the hour
                current_hour = datetime.now().hour
                if self.last_report_hour != current_hour:
                    print("Generating hourly report...")
                    self.make_report(df, stats)
                    self.last_report_hour = current_hour
                
                print(f"Done! Tracking {len(df)} coins")
                print(f"Total market cap: ${stats['market_overview']['total_market_cap']:,.2f}")
                
            except Exception as e:
                print(f"Something went wrong: {e}")
            
            # Wait before next update
            time.sleep(self.update_interval)

    # Start tracking
    def start(self):
        self.thread = threading.Thread(target=self.run)
        self.thread.start()
        print("Tracker started! Press Ctrl+C to stop")

    # Stop tracking
    def stop(self):
        self.running = False
        if hasattr(self, 'thread'):
            self.thread.join()
        print("Tracker stopped")


if __name__ == "__main__":
    tracker = CryptoTracker(update_interval=300)  # Update every 5 min
    
    try:
        tracker.start()
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\nStopping...")
        tracker.stop()