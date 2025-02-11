from flask import Flask, render_template_string
import requests
import pandas as pd
import time
from datetime import datetime
import threading
from openpyxl.styles import PatternFill, Font
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns

app = Flask(__name__)

class CryptoWebTracker:
    def __init__(self, update_interval=300):
        self.endpoints = [
            "https://api.binance.com",
            "https://api-gcp.binance.com",
            "https://api1.binance.com",
        ]
        self.data_api = "https://data-api.binance.vision"
        self.coingecko = "https://api.coingecko.com/api/v3"
        self.update_interval = update_interval
        self.last_report_hour = None
        
        # Output files
        self.excel_file = "crypto_data.xlsx"
        self.report_file = "market_report.docx"
        
        # Web data storage
        self.latest_data = {
            'data': None,
            'stats': None,
            'last_updated': None
        }
        
        self.running = True

    def get_coin_names(self):
        try:
            r = requests.get(f"{self.coingecko}/coins/list")
            
            if r.status_code == 200:
                coins = r.json()
                return {c['symbol'].upper(): c['name'] for c in coins}
            
            print("Couldn't fetch coin names")
            return {}
            
        except Exception as e:
            print(f"Error getting coin names: {e}")
            return {}

    def get_market_data(self):
        try:
            r = requests.get(f"{self.data_api}/api/v3/ticker/24hr")
            
            if r.status_code != 200:
                print("API call failed")
                return None
                
            data = r.json()
            
            usdt_pairs = [x for x in data if x['symbol'].endswith('USDT')]
            
            names = self.get_coin_names()
            for pair in usdt_pairs:
                sym = pair['symbol'].replace('USDT', '').upper()
                pair['name'] = names.get(sym, sym)  # Use symbol if name not found
            
            return usdt_pairs
            
        except Exception as e:
            print(f"Failed to get market data: {e}")
            return None

    def process_data(self, raw_data):
        df = pd.DataFrame(raw_data)
        
        # Convert strings to numbers
        number_cols = [
            'lastPrice', 'volume', 'quoteVolume', 
            'priceChangePercent', 'weightedAvgPrice'
        ]
        
        for col in number_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # Sort by volume and get top 50
        df = df.sort_values('quoteVolume', ascending=False).head(50)
        
        # Calculate market cap (price * volume)
        df['market_cap'] = df['lastPrice'] * df['volume']
        
        final = pd.DataFrame({
            'Name': df['name'],
            'Symbol': df['symbol'].str.replace('USDT', ''),
            'Price': df['lastPrice'],
            'Market Cap': df['market_cap'],
            'Volume (24h)': df['volume'],
            'Change (24h)': df['priceChangePercent'],
            'Avg Price': df['weightedAvgPrice'],
        })
        
        return final

    def make_chart(self, df):
        plt.figure(figsize=(12, 6))
        
        # Distribution of 24h price changes
        sns.histplot(data=df, x='Change (24h)', bins=20)
        plt.title('24h Price Changes')
        plt.xlabel('Change (%)')
        plt.ylabel('Number of Coins')
        
        plt.savefig('price_changes.png')
        plt.close()

    def analyze_market(self, df):
        # Calculate market overview statistics
        stats = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'market_overview': {
                'total_market_cap': df['Market Cap'].sum(),
                'total_volume_24h': df['Volume (24h)'].sum(),
                'avg_price': df['Price'].mean(),
                'median_price': df['Price'].median(),
            },
            'market_health': {
                'positive_performers': len(df[df['Change (24h)'] > 0]),
                'negative_performers': len(df[df['Change (24h)'] < 0]),
            },
            'detailed_analysis': {
                'top_5_by_market_cap': df.nlargest(5, 'Market Cap')[['Name', 'Symbol', 'Market Cap', 'Price']].to_dict('records'),
                'price_statistics': {
                    'average_price': df['Price'].mean(),
                    'median_price': df['Price'].median(),
                    'highest_price': df['Price'].max(),
                    'lowest_price': df['Price'].min(),
                },
                'price_changes': {
                    'highest_24h_change': df.nlargest(1, 'Change (24h)')[['Name', 'Symbol', 'Change (24h)']].to_dict('records')[0],
                    'lowest_24h_change': df.nsmallest(1, 'Change (24h)')[['Name', 'Symbol', 'Change (24h)']].to_dict('records')[0],
                    'average_24h_change': df['Change (24h)'].mean(),
                }
            }
        }
        return stats

    def make_report(self, df, stats):
        doc = Document()
        
        # Title
        doc.add_heading('Crypto Market Report', 0)
        doc.add_paragraph(f"Generated: {stats['timestamp']}")
        
        # Market overview section
        doc.add_heading('Market Overview', 1)
        overview = stats['market_overview']
        doc.add_paragraph(f"Total Market Cap: ${overview['total_market_cap']:,.2f}")
        doc.add_paragraph(f"24h Volume: ${overview['total_volume_24h']:,.2f}")
        
        # Top coins section with more details
        doc.add_heading('Top 5 Cryptocurrencies by Market Cap', 1)
        for coin in stats['detailed_analysis']['top_5_by_market_cap']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{coin['Name']} ({coin['Symbol']})\n").bold = True
            p.add_run(f"Market Cap: ${coin['Market Cap']:,.2f}\n")
            p.add_run(f"Price: ${coin['Price']:,.2f}")
        
        # Price Statistics
        doc.add_heading('Price Statistics', 1)
        price_stats = stats['detailed_analysis']['price_statistics']
        doc.add_paragraph(f"Average Price: ${price_stats['average_price']:,.2f}")
        doc.add_paragraph(f"Median Price: ${price_stats['median_price']:,.2f}")
        doc.add_paragraph(f"Highest Price: ${price_stats['highest_price']:,.2f}")
        doc.add_paragraph(f"Lowest Price: ${price_stats['lowest_price']:,.2f}")
        
        # 24h Price Changes
        doc.add_heading('24-Hour Price Changes', 1)
        changes = stats['detailed_analysis']['price_changes']
        best = changes['highest_24h_change']
        worst = changes['lowest_24h_change']
        
        doc.add_paragraph(f"Highest Gainer: {best['Name']} ({best['Symbol']}) with {best['Change (24h)']:+.2f}%")
        doc.add_paragraph(f"Biggest Decliner: {worst['Name']} ({worst['Symbol']}) with {worst['Change (24h)']:+.2f}%")
        doc.add_paragraph(f"Average 24h Change: {changes['average_24h_change']:+.2f}%")
        
        # Market health section
        doc.add_heading('Market Health', 1)
        health = stats['market_health']
        doc.add_paragraph(f"Coins Up: {health['positive_performers']}")
        doc.add_paragraph(f"Coins Down: {health['negative_performers']}")
        
        # Add a price change chart
        self.make_chart(df)
        doc.add_picture('price_changes.png', width=Inches(6))
        
        doc.save(self.report_file)

    def update_excel(self, df, stats):
        with pd.ExcelWriter(self.excel_file, engine='openpyxl') as writer:
            # Main data sheet
            df.to_excel(writer, sheet_name='Market Data', index=False)
            
            # Analysis sheet
            analysis = [
                ['Last Updated', stats['timestamp']],
                ['Market Overview', ''],
                ['Total Market Cap', f"${stats['market_overview']['total_market_cap']:,.2f}"],
                ['24h Volume', f"${stats['market_overview']['total_volume_24h']:,.2f}"],
                ['Average Price', f"${stats['market_overview']['avg_price']:,.2f}"],
                ['', ''],
                ['Market Health', ''],
                ['Coins Up', stats['market_health']['positive_performers']],
                ['Coins Down', stats['market_health']['negative_performers']],
            ]
            
            pd.DataFrame(analysis).to_excel(
                writer, sheet_name='Analysis',
                index=False, header=False
            )
            
            # Make it look nice
            self._format_excel(writer)

    # Add some color to the Excel sheets
    def _format_excel(self, writer):
        for sheet in ['Market Data', 'Analysis']:
            ws = writer.sheets[sheet]
            # Color the headers
            for cell in ws[1]:
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)

    def run(self):
        while self.running:
            print(f"\nUpdating data... {datetime.now().strftime('%H:%M:%S')}")
            
            try:
                # Get fresh data
                data = self.get_market_data()
                if not data:
                    continue
                    
                # Process it
                df = self.process_data(data)
                stats = self.analyze_market(df)
                
                # Update files
                self.update_excel(df, stats)
                
                # Update web data
                self.latest_data = {
                    'data': df.to_dict('records'),
                    'stats': stats,
                    'last_updated': stats['timestamp']
                }
                
                # Make report on the hour
                current_hour = datetime.now().hour
                if self.last_report_hour != current_hour:
                    print("Generating hourly report...")
                    self.make_report(df, stats)
                    self.last_report_hour = current_hour
                
                print(f"Done! Tracking {len(df)} coins")
                print(f"Total market cap: ${stats['market_overview']['total_market_cap']:,.2f}")
                
            except Exception as e:
                print(f"Something went wrong: {e}")
            
            time.sleep(self.update_interval)

tracker = CryptoWebTracker()

# HTML template
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Crypto Market Tracker</title>
    <style>
        :root {
            --primary: #2c3e50;
            --secondary: #34495e;
            --accent: #3498db;
            --success: #2ecc71;
            --danger: #e74c3c;
            --text: #2c3e50;
            --light-bg: #f8f9fa;
            --border: #dee2e6;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
            line-height: 1.6;
            color: var(--text);
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            min-height: 100vh;
        }

        .container {
            max-width: 1200px;
            margin: 2rem auto;
            padding: 0 1rem;
        }

        h1 {
            text-align: center;
            color: var(--primary);
            font-size: 2.5rem;
            margin-bottom: 1rem;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
        }

        .timestamp {
            text-align: center;
            color: var(--secondary);
            margin-bottom: 2rem;
            font-size: 0.9rem;
        }

        .refresh-info {
            background: white;
            padding: 1rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 2rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 1rem;
        }

        .refresh-button {
            background: var(--accent);
            color: white;
            border: none;
            padding: 0.5rem 1rem;
            border-radius: 4px;
            cursor: pointer;
            transition: transform 0.2s, background 0.2s;
            font-weight: 500;
            text-align: center;
        }

        .refresh-button:hover {
            background: #2980b9;
            transform: translateY(-1px);
        }

        .summary {
            background: white;
            padding: 1.5rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 2rem;
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1.5rem;
        }

        .summary h2 {
            grid-column: 1/-1;
            color: var(--primary);
            margin-bottom: 0.5rem;
        }

        .summary p {
            background: var(--light-bg);
            padding: 1rem;
            border-radius: 6px;
            text-align: center;
            font-weight: 500;
        }

        table {
            width: 100%;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            border-collapse: collapse;
            overflow: hidden;
        }

        thead {
            background: var(--primary);
            color: white;
        }

        th {
            padding: 1rem;
            text-align: left;
            font-weight: 600;
        }

        td {
            padding: 1rem;
            border-bottom: 1px solid var(--border);
        }

        tbody tr:hover {
            background: var(--light-bg);
        }

        .positive {
            color: var(--success);
            font-weight: 600;
        }

        .negative {
            color: var(--danger);
            font-weight: 600;
        }

        #loading-overlay {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(255,255,255,0.9);
            display: none;
            justify-content: center;
            align-items: center;
            z-index: 1000;
        }

        .loading-spinner {
            width: 50px;
            height: 50px;
            border: 5px solid var(--border);
            border-top: 5px solid var(--accent);
            border-radius: 50%;
            animation: spin 1s linear infinite;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        @media (max-width: 768px) {
            .container {
                margin: 1rem auto;
            }

            .refresh-info {
                flex-direction: column;
                text-align: center;
            }

            .refresh-button {
                width: 100%;
            }

            td, th {
                padding: 0.75rem 0.5rem;
                font-size: 0.9rem;
            }
        }
    </style>
</head>
<body>
    <div id="loading-overlay">
        <div class="loading-spinner"></div>
    </div>
    <div class="container">
        <h1>Crypto Market Tracker</h1>
        <p class="timestamp">Last Updated: {{ last_updated }}</p>
        
        <div class="refresh-info">
            <span>Next refresh in: <span id="countdown">300</span> seconds</span>
            <button class="refresh-button" onclick="refreshNow()">Refresh Now</button>
            <a href="/crypto_data.xlsx" class="refresh-button">Download Excel</a>
            <a href="/market_report.docx" class="refresh-button">Download Report</a>
        </div>

        <div class="summary">
            <h2>Market Summary</h2>
            <p>Average Price<br>${{ "{:,.2f}".format(stats['market_overview']['avg_price']) }}</p>
            <p>Total Market Cap<br>${{ "{:,.2f}".format(stats['market_overview']['total_market_cap']) }}</p>
            <p>24h Volume<br>${{ "{:,.2f}".format(stats['market_overview']['total_volume_24h']) }}</p>
            <p>Market Health<br>↑ {{ stats['market_health']['positive_performers'] }} | ↓ {{ stats['market_health']['negative_performers'] }}</p>
        </div>

        <table>
            <thead>
                <tr>
                    <th>Rank</th>
                    <th>Name</th>
                    <th>Symbol</th>
                    <th>Price (USD)</th>
                    <th>Market Cap</th>
                    <th>24h Change</th>
                    <th>24h Volume</th>
                </tr>
            </thead>
            <tbody>
                {% for coin in data %}
                <tr>
                    <td>{{ loop.index }}</td>
                    <td>{{ coin['Name'] }}</td>
                    <td>{{ coin['Symbol'] }}</td>
                    <td>${{ "{:,.2f}".format(coin['Price']) }}</td>
                    <td>${{ "{:,.2f}".format(coin['Market Cap']) }}</td>
                    <td class="{{ 'positive' if coin['Change (24h)'] > 0 else 'negative' }}">
                        {{ "{:+.2f}%".format(coin['Change (24h)']) }}
                    </td>
                    <td>${{ "{:,.2f}".format(coin['Volume (24h)']) }}</td>
                </tr>
                {% endfor %}
            </tbody>
        </table>
    </div>

    <script>
        let countdownValue = 300;
        const countdownElement = document.getElementById('countdown');
        
        function updateCountdown() {
            countdownValue--;
            countdownElement.textContent = countdownValue;
            
            if (countdownValue <= 0) {
                location.reload();
            }
        }
        
        setInterval(updateCountdown, 1000);
        
        function refreshNow() {
            document.getElementById('loading-overlay').style.display = 'flex';
            location.reload();
        }
    </script>
</body>
</html>
'''

@app.route('/')
def home():
    if tracker.latest_data['data'] is None:
        return "Loading data... Please refresh in a moment."
    
    return render_template_string(
        HTML_TEMPLATE,
        data=tracker.latest_data['data'],
        stats=tracker.latest_data['stats'],
        last_updated=tracker.latest_data['last_updated'],
        excel_path=tracker.excel_file,
        report_path=tracker.report_file
    )

def start_app():
    # Start the data tracking thread
    tracking_thread = threading.Thread(target=tracker.run)
    tracking_thread.daemon = True
    tracking_thread.start()
    
    app.run(host='0.0.0.0', port=5000)

if __name__ == '__main__':
    start_app()